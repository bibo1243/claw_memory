#!/usr/bin/env python3
"""產生 2026.02.26 慈光核心主管會議紀錄 Word 檔 — 使用 Word 內建自動大綱編號"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import OxmlElement, parse_xml
from lxml import etree
import copy

doc = Document()

# === 設定預設字型 ===
style = doc.styles['Normal']
font = style.font
font.name = '標楷體'
font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')
style.paragraph_format.line_spacing = 1.15

# === 頁面邊距 ===
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

# =============================================
# 建立自動大綱編號定義 (Multilevel List)
# 壹、→ 一、→（一）→ 1. → (1)
# =============================================

# 取得或建立 numbering part
numbering_part = doc.part.numbering_part
numbering_elm = numbering_part._element

# 找到最大的 abstractNumId
existing_abstract = numbering_elm.findall(qn('w:abstractNum'))
max_abstract_id = max([int(a.get(qn('w:abstractNumId'))) for a in existing_abstract], default=-1)
abstract_num_id = max_abstract_id + 1

existing_nums = numbering_elm.findall(qn('w:num'))
max_num_id = max([int(n.get(qn('w:numId'))) for n in existing_nums], default=0)
num_id = max_num_id + 1

# 定義五層大綱
# Level 0: 壹、貳、參、肆、伍、陸、柒
# Level 1: 一、二、三、四、五
# Level 2: （一）（二）（三）
# Level 3: 1. 2. 3.
# Level 4: (1) (2) (3)

levels_config = [
    {   # Level 0: 壹、
        'numFmt': 'ideographLegalTraditional',
        'lvlText': '%1、',
        'left': 720,    # 0.5 inch = 720 twips  
        'hanging': 480,
        'font_size': 15,
        'bold': True,
    },
    {   # Level 1: 一、
        'numFmt': 'taiwaneseCounting',
        'lvlText': '%2、',
        'left': 1200,
        'hanging': 480,
        'font_size': 13,
        'bold': True,
    },
    {   # Level 2: （一）
        'numFmt': 'taiwaneseCounting',
        'lvlText': '（%3）',
        'left': 1680,
        'hanging': 480,
        'font_size': 12,
        'bold': False,
    },
    {   # Level 3: 1.
        'numFmt': 'decimal',
        'lvlText': '%4.',
        'left': 2160,
        'hanging': 360,
        'font_size': 12,
        'bold': False,
    },
    {   # Level 4: (1)
        'numFmt': 'decimal',
        'lvlText': '(%5)',
        'left': 2640,
        'hanging': 360,
        'font_size': 12,
        'bold': False,
    },
]

# 建立 abstractNum XML
abstract_num_xml = f'<w:abstractNum w:abstractNumId="{abstract_num_id}" {nsdecls("w")}>'
abstract_num_xml += '<w:multiLevelType w:val="multilevel"/>'

for i, cfg in enumerate(levels_config):
    abstract_num_xml += f'''
    <w:lvl w:ilvl="{i}">
        <w:start w:val="1"/>
        <w:numFmt w:val="{cfg['numFmt']}"/>
        <w:suff w:val="nothing"/>
        <w:lvlText w:val="{cfg['lvlText']}"/>
        <w:lvlJc w:val="left"/>
        <w:pPr>
            <w:ind w:left="{cfg['left']}" w:hanging="{cfg['hanging']}"/>
        </w:pPr>
        <w:rPr>
            <w:rFonts w:ascii="標楷體" w:eastAsia="標楷體" w:hAnsi="標楷體"/>
            <w:sz w:val="{cfg['font_size'] * 2}"/>
            <w:b w:val="{'1' if cfg['bold'] else '0'}"/>
        </w:rPr>
    </w:lvl>'''

abstract_num_xml += '</w:abstractNum>'

abstract_num_element = parse_xml(abstract_num_xml)
# 插入到最後一個 abstractNum 之後
if existing_abstract:
    existing_abstract[-1].addnext(abstract_num_element)
else:
    numbering_elm.append(abstract_num_element)

# 建立 num 引用
num_xml = f'<w:num w:numId="{num_id}" {nsdecls("w")}><w:abstractNumId w:val="{abstract_num_id}"/></w:num>'
num_element = parse_xml(num_xml)
numbering_elm.append(num_element)

# =============================================
# Helper 函式
# =============================================

def add_outline_item(text, level, bold=None, restart=False):
    """使用自動大綱編號新增段落"""
    p = doc.add_paragraph()
    pPr = p._element.get_or_add_pPr()
    
    # 設定 numbering
    numPr = OxmlElement('w:numPr')
    ilvl = OxmlElement('w:ilvl')
    ilvl.set(qn('w:val'), str(level))
    numPr.append(ilvl)
    numId_elm = OxmlElement('w:numId')
    numId_elm.set(qn('w:val'), str(num_id))
    numPr.append(numId_elm)
    pPr.append(numPr)
    
    # 行距（緊湊）
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), '0')
    spacing.set(qn('w:after'), '0')
    spacing.set(qn('w:line'), '276')  # 1.15倍行距 (240*1.15=276)
    spacing.set(qn('w:lineRule'), 'auto')
    pPr.append(spacing)
    
    # 加入文字
    cfg = levels_config[level]
    is_bold = bold if bold is not None else cfg['bold']
    
    run = p.add_run(text)
    run.font.name = '標楷體'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')
    run.font.size = Pt(cfg['font_size'])
    run.bold = is_bold
    
    return p

def add_plain_para(text, bold=False, font_size=12, alignment=None, indent_cm=0):
    """新增不帶編號的普通段落"""
    p = doc.add_paragraph()
    if alignment:
        p.alignment = alignment
    if indent_cm > 0:
        p.paragraph_format.left_indent = Cm(indent_cm)
    run = p.add_run(text)
    run.font.name = '標楷體'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')
    run.font.size = Pt(font_size)
    run.bold = bold
    return p

def reset_numbering_at_level(level):
    """重設某層級的編號（從1重新開始）"""
    # 建立一個新的 num 指向同一個 abstractNum 但有 lvlOverride
    global num_id
    max_nums = numbering_elm.findall(qn('w:num'))
    new_num_id = max([int(n.get(qn('w:numId'))) for n in max_nums], default=0) + 1
    
    override_xml = f'<w:num w:numId="{new_num_id}" {nsdecls("w")}>'
    override_xml += f'<w:abstractNumId w:val="{abstract_num_id}"/>'
    override_xml += f'<w:lvlOverride w:ilvl="{level}">'
    override_xml += '<w:startOverride w:val="1"/>'
    override_xml += '</w:lvlOverride>'
    override_xml += '</w:num>'
    
    num_element = parse_xml(override_xml)
    numbering_elm.append(num_element)
    num_id = new_num_id
    return new_num_id

# =============================================
# 文件標題區
# =============================================
add_plain_para('財團法人台中市私立慈光社會福利慈善事業基金會', bold=True, font_size=16, alignment=WD_ALIGN_PARAGRAPH.CENTER)
add_plain_para('慈光核心主管會議紀錄', bold=True, font_size=16, alignment=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()

# 會議資訊 — 無框線純文字段落
def add_info_line(label, value):
    p = doc.add_paragraph()
    pPr = p._element.get_or_add_pPr()
    # 段落間距緊湊
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), '0')
    spacing.set(qn('w:after'), '0')
    spacing.set(qn('w:line'), '276')
    spacing.set(qn('w:lineRule'), 'auto')
    pPr.append(spacing)
    # Label（粗體）
    run_label = p.add_run(label)
    run_label.font.name = '標楷體'
    run_label._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')
    run_label.font.size = Pt(12)
    run_label.bold = True
    # Value（緊接在冒號後）
    run_value = p.add_run(value)
    run_value.font.name = '標楷體'
    run_value._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')
    run_value.font.size = Pt(12)
    return p

add_info_line('日　期：', '115年02月26日（四）')
add_info_line('時　間：', '上午9:00-12:00')
add_info_line('地　點：', '基金會會議室')
add_info_line('主　席：', '詹前柏常務董事')
add_info_line('出　席：', '吳碧霞董事長、廖慧雯主任、李冠葦總幹事')
add_info_line('紀　錄：', '李冠葦總幹事')

doc.add_paragraph()

# =============================================
# 壹、追蹤事項
# =============================================
add_outline_item('追蹤事項', level=0)

add_outline_item('上次會議事項追蹤', level=1)

add_outline_item('勞動檢查後續處理', level=2, bold=True)
add_outline_item('目前已致電勞工局約五次，承辦人不在，有留電話但未回電。', level=3)
add_outline_item('行政裁處目前取消。', level=3)
add_outline_item('基金會已朝合法勞動基準方向調整，基本上已合法。', level=3)
add_outline_item('聯合稽查勞檢通常在一至兩個月內進行，屬常態時程。', level=3)

add_outline_item('值班間隔時數問題', level=2, bold=True)
reset_numbering_at_level(3)
add_outline_item('討論值晚班人員至隔日上班之間隔時數是否符合規定。', level=3)
add_outline_item('輪班制需間隔11小時；本機構經確認非屬輪班制。', level=3)
add_outline_item('決議：請冠葦致電勞工局諮詢確認加班後與隔日上班之間隔規定，朝間隔8小時以上方向處理。', level=3)

# =============================================
# 貳、報告及討論事項——慈光基金會
# =============================================
add_outline_item('報告及討論事項——慈光基金會', level=0)

# === 議題一 ===
add_outline_item('員工育兒支持方案實施辦法修訂（附件一）', level=1)

add_outline_item('提案說明', level=2, bold=True)
reset_numbering_at_level(3)
add_outline_item('振杉主任反映：育兒津貼是否應納入平均工資計算，影響加班費、月薪及時薪。', level=3)
add_outline_item('常務董事認為：育兒津貼屬員工福利措施，非工作對價之給付，不具經常性工資之性質，不應列入平均工資、加班費、退休金等計算基礎。', level=3)

add_outline_item('討論重點', level=2, bold=True)
reset_numbering_at_level(3)
add_outline_item('國稅局認定：凡有所得皆須繳稅（薪資所得），此為必然，與是否列入平均工資係屬不同層次問題。', level=3)
add_outline_item('育兒津貼之發放方式調整：不再併入薪資表，改以支出憑證方式另行核銷，與薪資分開匯款。', level=3)
add_outline_item('久任獎金、育兒津貼、助學金均統一由梅芳每月製作支出憑證，送交各機構主管用印後核銷。', level=3)
add_outline_item('留職停薪期間，相關福利（含育兒津貼、助學金、考績獎金、年終獎金等）一律暫停。建議於工作規則中統一規範，不需逐一辦法標註。', level=3)
add_outline_item('未來組織所有辦法應統一做法：以工作規則為母法，留職停薪期間所有福利暫停。', level=3)

add_outline_item('決議', level=2, bold=True)
reset_numbering_at_level(3)
add_outline_item('辦法增列一條說明：本方案所提供之育兒津貼及助學金屬員工福利措施，非勞務對價之給付，不列入平均工資等計算基礎。', level=3)
add_outline_item('育兒津貼、助學金及久任獎金之發放改以支出憑證方式辦理，由梅芳統一製作、送機構主管用印後核撥。', level=3)
add_outline_item('三月份會務會議向同仁布達後實施。', level=3)
add_outline_item('董事長同意修訂，無須提請董事會。', level=3)

# === 議題二 ===
add_outline_item('董事會流程之規劃討論', level=1)

add_outline_item('議程討論（附件二）', level=2, bold=True)
reset_numbering_at_level(3)
add_outline_item('基本議程與往年相同：業務執行報告、收支餘絀表暨資產負債表、財產清冊、各基金決算表、受贈財物使用情形、內部控制制度檢核等。', level=3)
add_outline_item('原議程第十四案「會計制度修訂」取消（無修訂事項）。', level=3)
add_outline_item('原議程第十六案「流用基金管理辦法」改於機構主管會議確認即可，無須提董事會（原訂定時為主任公布施行，未經董事會）。', level=3)
add_outline_item('兒少之家短中長程發展計畫執行報告仍需提報。', level=3)

add_outline_item('董事會開會時間', level=2, bold=True)
reset_numbering_at_level(3)
add_outline_item('吳董事長4/26下午有行程，無法與會。', level=3)
add_outline_item('決議：維持4/26上午開會，以吳董事長時間為主。', level=3)
add_outline_item('李常董（李宏偉老師）因在台北，時間較為困難；若能來則參加，不勉強。需向吳董事長報告此情形，並回覆李常董。', level=3)

add_outline_item('兒少之家建物借款之會計處理', level=2, bold=True)
reset_numbering_at_level(3)
add_outline_item('會計師李會計師來電詢問：基金會蓋少年家園時向兒少之家借款6,400萬元，記為「作業組織往來」科目，是否需歸還？', level=3)
add_outline_item('若不還：兒少之家營運資金可能呈負數，且涉及責信問題（評鑑時會計師可能有意見）。', level=3)
add_outline_item('討論：母會（基金會）撥款給子機構在法理上可行，但反向可能有問題。', level=3)
add_outline_item('決議：維持「作業組織往來」，基金會將來仍需歸還。此為既定方向，無須另提董事會議案。未來基金會捐款逐漸增加後，應逐步償還。', level=3)

add_outline_item('董事資料閱覽方式', level=2, bold=True)
reset_numbering_at_level(3)
add_outline_item('陳董事反映希望事前閱覽決算資料。', level=3)
add_outline_item('處理方式：業務執行報告書事前送交董事參考；決算暫不事前寄送，可約時間至現場閱覽後收回。', level=3)

add_outline_item('報告事項提醒', level=2, bold=True)
reset_numbering_at_level(3)
add_outline_item('報告時須向董事提及明年有評鑑事宜。', level=3)

# === 議題三 ===
add_outline_item('家計經營計畫未來形式討論（附件三：專業顧問會議摘要紀錄）', level=1)

add_outline_item('背景說明', level=2, bold=True)
reset_numbering_at_level(3)
add_outline_item('上次顧問會議已討論家計經營計畫之存廢與轉型，董事長當時裁示暫不決議，大家繼續思考。', level=3)
add_outline_item('本次會議延續討論未來方向。', level=3)

add_outline_item('現場狀況與困境', level=2, bold=True)
reset_numbering_at_level(3)
add_outline_item('保育/生輔人員撰寫經驗整理對工作負擔極大，須使用工作以外時間，督導負擔亦重。', level=3)
add_outline_item('現場照顧工作繁重，尤其晚間意外事件多，人力已相當吃緊。', level=3)
add_outline_item('目前輪班制度下，「小家經營」概念已與過去不同。', level=3)

add_outline_item('各方意見', level=2, bold=True)
reset_numbering_at_level(3)
add_outline_item('常務董事：支持轉型，至少不再稱「家計經營計畫」。工作經驗整理有價值，但需有配套措施（獎勵、加班費計算等）。', level=3)
add_outline_item('董事長：家計經營計畫之功能應回到「記日記般的經驗整理」，對同仁個人生命有幫助，重點是鼓勵而非強制。以獎勵作為動力。', level=3)
add_outline_item('廖主任：工作型態已不同，現場同仁對孩子是真心付出但專業度待提升。經驗整理產出的效益是存在的。建議關注組織使命與經營哲學的釐清。', level=3)

add_outline_item('決議', level=2, bold=True)
reset_numbering_at_level(3)
add_outline_item('本議題暫不做最終決定，持續研議。', level=3)
add_outline_item('需與振杉主任進一步討論，聽取現場教保團隊意見。', level=3)
add_outline_item('三位專業顧問繼續參與後續思考。', level=3)
add_outline_item('下一步需回答之核心問題：', level=3)
add_outline_item('組織下一個5-10年的使命和經營哲學是什麼？', level=4)
add_outline_item('要養出什麼樣的孩子？創造什麼樣的氛圍？', level=4)
add_outline_item('如何導入AI輔助，減輕行政與紀錄負擔？', level=4)
add_outline_item('激勵制度如何設計？', level=4)
add_outline_item('人力問題如何解決？經費從哪裡來？', level=4)

# =============================================
# 參、行政管理事項
# =============================================
add_outline_item('行政管理事項', level=0)

add_outline_item('員工在職訓練體系討論', level=1)

add_outline_item('職前訓練', level=2, bold=True)
reset_numbering_at_level(3)
add_outline_item('法規規定專業人員職前訓練6小時、在職訓練每年18小時。', level=3)
add_outline_item('廖主任建議可統一規劃6小時職前訓練架構：工作手冊介紹（1小時）、組織歷史與發展簡報（1-2小時）、本務職務說明、工作規則說明。', level=3)
add_outline_item('期待人力資源暨公共關係組（梅芳）能發展此系統，惟考量其能力限制，可善用線上課程、錄影等資源。', level=3)
add_outline_item('振杉主任在執行面是否有落實新進人員職前訓練，需再確認。', level=3)

add_outline_item('法定課程與訓練管制', level=2, bold=True)
reset_numbering_at_level(3)
add_outline_item('四項法定課程：衛生教育訓練（今年規定全體員工4小時）、事故訓練、CRC兒童權利公約訓練（今年擴增至3小時）、性侵害防治。', level=3)
add_outline_item('訓練要求基本態度為百分之百出席，無法出席須外部補課並提供證書。', level=3)
add_outline_item('今年訓練與會議統籌安排由廖主任帶領淑錡進行。基金會上半年課程已排出。', level=3)

add_outline_item('專業訓練內容檢討', level=2, bold=True)
reset_numbering_at_level(3)
add_outline_item('廖主任指出：兒家訓練課程偏重情感支援與自我覺察，缺少安置兒少類型認識與實務提升課程。', level=3)
add_outline_item('瑞華老師（員工支援團體）之角色定位：以團體支持、人際靠近為主，非專業訓練。同仁反應良好，建議維持一季一次。', level=3)
add_outline_item('專業訓練部分改請顗帆（諮商所）協助規劃，依董事會去年決議，諮商所應承擔組織內部專業訓練之義務。', level=3)
add_outline_item('CRC訓練今年與具機構實務經驗之講師長期合作，擴增至3小時含分組討論。', level=3)

add_outline_item('行政辦公室管理議題', level=1)

add_outline_item('辦公室動力與同仁態度', level=2, bold=True)
reset_numbering_at_level(3)
add_outline_item('行政人員對許多事情「不清不楚」，傳遞訊息變少、疑惑增多。', level=3)
add_outline_item('淑錡反映辦公室動力問題，包含：來賓接待意識薄弱、對孩子狀況不了解導致合作困難、八卦傳言等。', level=3)
add_outline_item('常務董事提醒：一步一步慢慢改善，不可能一下子到位。振杉可適當協助。', level=3)

add_outline_item('淑錡（副組長）管理風格', level=2, bold=True)
reset_numbering_at_level(3)
add_outline_item('廖主任觀察：淑錡對同仁講話態度偏權威、有時過於嚴厲，與組織尊重員工的文化不符。', level=3)
add_outline_item('建議振杉與淑錡建立「事先約定」機制：約定好在辦公室中若出現不當態度，可直接提醒，作為兩位主管的示範。', level=3)

add_outline_item('督導機制', level=2, bold=True)
reset_numbering_at_level(3)
add_outline_item('淑錡與冠葦之間的督導配合仍需加強。常務董事提醒：重視度不夠、資訊給予太晚。', level=3)
add_outline_item('理想做法：前兩天給資訊，會前會一小時、督導後會後會再一小時。', level=3)

add_outline_item('機構空間盤點與規劃', level=1)

add_outline_item('問題現況', level=2, bold=True)
reset_numbering_at_level(3)
add_outline_item('行政辦公室桌子太小、東西堆積過多。', level=3)
add_outline_item('淑錡書桌下方有空間但因壁虎問題未使用。', level=3)
add_outline_item('梅芳佔用一間獨立辦公室。', level=3)

add_outline_item('可盤點空間清單', level=2, bold=True)
reset_numbering_at_level(3)
add_outline_item('董事長辦公位後方木櫃（有一區已撤出可用）。', level=3)
add_outline_item('少家一樓櫃子（原社資組使用，已移出）。', level=3)
add_outline_item('少家三樓公佈欄下方櫃子（原衛生保健庫存）。', level=3)
add_outline_item('員工宿舍五樓下方櫃子。', level=3)
add_outline_item('少家七樓櫃子（籃球場旁）。', level=3)
add_outline_item('小家男一、男二、男三閣樓（已回收清空上鎖）。', level=3)
add_outline_item('少家一樓茶水間下方空間。', level=3)
add_outline_item('二樓諮商室外房間（原寢房）。', level=3)

add_outline_item('原則', level=2, bold=True)
reset_numbering_at_level(3)
add_outline_item('兩棟建築空間皆應做到「寸土寸金」最大化利用。', level=3)
add_outline_item('凡使用少家空間，須讓廖主任知悉，以明確責任歸屬。', level=3)
add_outline_item('應做全機構空間整體盤點與規劃。', level=3)

add_outline_item('人事管理事項', level=1)

add_outline_item('公文與對外文件管理', level=2, bold=True)
reset_numbering_at_level(3)
add_outline_item('所有對外發出之檔案、公文，必須讓機構主管知悉。', level=3)
add_outline_item('具體案例：鳳翎代辦員工職災理賠申請（需蓋負責人章），未經機構主管審閱即送出，需改善。', level=3)
add_outline_item('人事相關文件尤須回歸機構主管把關，涉及保險理賠更需謹慎。', level=3)

add_outline_item('工作規則宣導', level=2, bold=True)
reset_numbering_at_level(3)
add_outline_item('工作規則為組織母法，同仁應充分了解。', level=3)
add_outline_item('廖主任從去年開始，新進同仁到任前即帶著工作規則逐條說明。', level=3)
add_outline_item('未來應在會務會議上定期進行工作規則相關宣導。', level=3)

# =============================================
# 肆、其他事項
# =============================================
add_outline_item('其他事項', level=0)

add_outline_item('生活輔導員培訓與證照：社家署正推動生活輔導員培訓並核發證照，台中7月開班，可考慮招募二度就業媽媽（高中以上學歷即可報名）。', level=1)
add_outline_item('現場人力觀察：新進同仁需有一定成熟度（約30歲以上較理想），不宜過早投入高強度照顧現場。可先從行政職切入，逐步接觸現場。', level=1)
add_outline_item('安置盟六大目標相關：人口結構問題、優化照顧模式、重新論述專業價值、發展多元替代照顧模式等。', level=1)

# =============================================
# 伍、待辦事項
# =============================================
add_outline_item('待辦事項', level=0)

add_plain_para('')

# 待辦事項表格
table = doc.add_table(rows=1, cols=4)
table.style = 'Table Grid'
headers = ['項次', '待辦事項', '負責人', '期限']
for i, header in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(header)
    run.bold = True
    run.font.name = '標楷體'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')
    run.font.size = Pt(11)

# Set column widths
for row in table.rows:
    row.cells[0].width = Cm(1.5)
    row.cells[1].width = Cm(10)
    row.cells[2].width = Cm(2.5)
    row.cells[3].width = Cm(3)

todo_items = [
    ('1', '育兒支持方案辦法增列非工資性質條款，提請董事長同意', '冠葦/梅芳', '3月會務會議前'),
    ('2', '育兒津貼、助學金、久任獎金改以支出憑證方式核銷', '梅芳', '3月份起'),
    ('3', '會務會議布達育兒津貼發放方式調整', '冠葦', '3月份會務會議'),
    ('4', '致電勞工局確認加班後與隔日上班間隔時數規定', '冠葦', '近期'),
    ('5', '董事會議程定稿（刪除第14、16案）', '冠葦', '董事會前'),
    ('6', '回覆李常董董事會時間安排', '冠葦', '近期'),
    ('7', '業務執行報告書送交董事參考', '冠葦', '董事會前'),
    ('8', '約陳董事閱覽決算資料時間', '冠葦', '董事會前'),
    ('9', '工作規則納入留職停薪期間福利暫停之統一規範', '梅芳', '配合修訂'),
    ('10', '釐清淑錡與鳳翎工作權責劃分', '冠葦/振杉', '近期'),
    ('11', '振杉與淑錡建立「行為約定」機制', '振杉', '近期'),
    ('12', '全機構空間盤點與規劃', '冠葦/振杉/廖主任', '本季'),
    ('13', '鳳翎對外文件須經機構主管審閱之程序建立', '冠葦', '即刻執行'),
    ('14', '家計經營計畫轉型方向持續研議', '常務董事/冠葦', '持續'),
]

for item in todo_items:
    row = table.add_row()
    for j, val in enumerate(item):
        cell = row.cells[j]
        cell.text = ''
        p = cell.paragraphs[0]
        if j == 0:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(val)
        run.font.name = '標楷體'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')
        run.font.size = Pt(11)

doc.add_paragraph()

# =============================================
# 陸、散會
# =============================================
add_outline_item('散會', level=0)
add_plain_para('散會時間：11:50', indent_cm=1)
add_plain_para('會後行程：12:00出發；18:00歲末圍爐（少年家園七樓多功能活動場）', indent_cm=1)

# === 保存 ===
output_path = '/Users/leegary/個人app/Inbox/2026.02.26慈光核心主管會議紀錄.docx'
doc.save(output_path)
print(f"✅ 會議紀錄已產出（使用Word自動大綱編號）：{output_path}")
