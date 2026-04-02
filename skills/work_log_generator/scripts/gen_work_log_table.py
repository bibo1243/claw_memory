#!/usr/bin/env python3
"""
工作日誌產出器 — 表格式 Word 文件
從 Markdown 草稿解析內容，產出三欄表格式工作日誌 (.docx)。

使用方式:
    python3 gen_work_log_table.py <草稿路徑> [輸出路徑]

範例:
    python3 gen_work_log_table.py ~/個人app/2026-03-06_WorkLog_Draft.md
    python3 gen_work_log_table.py ~/個人app/2026-03-06_WorkLog_Draft.md ~/個人app/個人資料夾/工作日誌/2026.03.06_工作日誌.docx
"""

import re
import sys
import os
from datetime import datetime, timedelta
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import OxmlElement, parse_xml
import docx.opc.constants


# =============================================
# 多層自動大綱編號引擎
# 六層架構：一、→（一）→ 1. →（1）→ a. →（a）
# =============================================

# 全域變數（在 generate_docx 中初始化）
_numbering_elm = None
_abstract_num_id = None
_current_num_id = None


def _build_outline_abstract_num_xml(abstract_num_id):
    """建立六層大綱編號的 abstractNum XML。
    
    Level 0: 一、二、三…  (taiwaneseCounting)
    Level 1: （一）（二）… (taiwaneseCounting + 括號)
    Level 2: 1. 2. 3.…    (decimal)
    Level 3: （1）（2）…   (decimal + 括號)
    Level 4: a. b. c.…    (lowerLetter)
    Level 5: (a) (b) (c)… (lowerLetter + 括號)
    """
    # 中文數字對應表
    cn_nums = '一二三四五六七八九十'
    
    levels = [
        # (numFmt, lvlText, left_twips, hanging_twips, font_size_half_pt, bold)
        ('taiwaneseCounting', '%1、',     480, 480, 24, False),  # 一、（取消粗體）
        ('taiwaneseCounting', '（%2）',   960, 480, 24, False),  # （一）
        ('decimal',           '%3.',     1440, 360, 24, False),  # 1.
        ('decimal',           '（%4）',  1800, 360, 24, False),  # （1）
        ('lowerLetter',       '%5.',     2160, 360, 24, False),  # a.
        ('lowerLetter',       '(%6)',    2520, 360, 24, False),  # (a)
    ]
    
    xml = f'<w:abstractNum w:abstractNumId="{abstract_num_id}" {nsdecls("w")}>'
    xml += '<w:multiLevelType w:val="multilevel"/>'
    
    for i, (num_fmt, lvl_text, left, hanging, sz, bold) in enumerate(levels):
        xml += f'<w:lvl w:ilvl="{i}">'
        xml += f'<w:start w:val="1"/>'
        xml += f'<w:numFmt w:val="{num_fmt}"/>'
        xml += f'<w:lvlText w:val="{lvl_text}"/>'
        xml += '<w:suff w:val="nothing"/>'
        xml += '<w:lvlJc w:val="left"/>'
        xml += f'<w:pPr><w:ind w:left="{left}" w:hanging="{hanging}"/></w:pPr>'
        xml += '<w:rPr>'
        xml += '<w:rFonts w:ascii="標楷體" w:eastAsia="標楷體" w:hAnsi="標楷體"/>'
        xml += f'<w:sz w:val="{sz}"/>'
        xml += f'<w:szCs w:val="{sz}"/>'
        if bold:
            xml += '<w:b/>'
        xml += '</w:rPr>'
        xml += '</w:lvl>'
    
    xml += '</w:abstractNum>'
    return xml


def init_outline_numbering(doc):
    """在文件中初始化多層大綱編號。必須在建表前呼叫。"""
    global _numbering_elm, _abstract_num_id, _current_num_id
    
    numbering_part = doc.part.numbering_part
    _numbering_elm = numbering_part._element
    
    # 找到可用的 abstractNumId
    existing = _numbering_elm.findall(qn('w:abstractNum'))
    _abstract_num_id = max([int(a.get(qn('w:abstractNumId'))) for a in existing], default=0) + 1
    
    # 建立 abstractNum
    abstract_xml = _build_outline_abstract_num_xml(_abstract_num_id)
    _numbering_elm.append(parse_xml(abstract_xml))
    
    # 建立初始 num
    existing_nums = _numbering_elm.findall(qn('w:num'))
    _current_num_id = max([int(n.get(qn('w:numId'))) for n in existing_nums], default=0) + 1
    
    num_xml = (
        f'<w:num w:numId="{_current_num_id}" {nsdecls("w")}>'
        f'<w:abstractNumId w:val="{_abstract_num_id}"/>'
        f'</w:num>'
    )
    _numbering_elm.append(parse_xml(num_xml))


def reset_outline_at_level(level):
    """重設指定層級的編號從 1 開始（例如切換到新的（一）時，讓 1. 重新計數）。"""
    global _current_num_id

    existing_nums = _numbering_elm.findall(qn('w:num'))
    _current_num_id = max([int(n.get(qn('w:numId'))) for n in existing_nums], default=0) + 1

    override_xml = (
        f'<w:num w:numId="{_current_num_id}" {nsdecls("w")}>'
        f'<w:abstractNumId w:val="{_abstract_num_id}"/>'
        f'<w:lvlOverride w:ilvl="{level}">'
        f'<w:startOverride w:val="1"/>'
        f'</w:lvlOverride>'
        f'</w:num>'
    )
    _numbering_elm.append(parse_xml(override_xml))


def _new_outline_num_id():
    """為每個 item cell 建立獨立的 numId，讓子層級都能從 1 重新開始。

    呼叫此函式後，level 0 ~ 5 都會從 1 重新計數，
    確保每個 add_data_row 的內容都從「一、」開始。
    """
    global _current_num_id

    existing_nums = _numbering_elm.findall(qn('w:num'))
    _current_num_id = max([int(n.get(qn('w:numId'))) for n in existing_nums], default=0) + 1

    # 同時對所有 6 層設 startOverride=1
    override_all = ''
    for lvl in range(6):
        override_all += (
            f'<w:lvlOverride w:ilvl="{lvl}">'
            f'<w:startOverride w:val="1"/>'
            f'</w:lvlOverride>'
        )

    num_xml = (
        f'<w:num w:numId="{_current_num_id}" {nsdecls("w")}>'
        f'<w:abstractNumId w:val="{_abstract_num_id}"/>'
        f'{override_all}'
        f'</w:num>'
    )
    _numbering_elm.append(parse_xml(num_xml))


def apply_numbering_to_paragraph(paragraph, level):
    """將自動大綱編號套用到段落上。
    
    level 0 = 一、
    level 1 = （一）
    level 2 = 1.
    level 3 = （1）
    level 4 = a.
    level 5 = (a)
    """
    pPr = paragraph._element.get_or_add_pPr()
    numPr = OxmlElement('w:numPr')
    ilvl = OxmlElement('w:ilvl')
    ilvl.set(qn('w:val'), str(level))
    numId = OxmlElement('w:numId')
    numId.set(qn('w:val'), str(_current_num_id))
    numPr.append(ilvl)
    numPr.append(numId)
    pPr.append(numPr)


# =============================================
# 解析 Markdown 草稿
# =============================================

def parse_markdown_links(text):
    """解析 Markdown 連結語法，回傳 [(文字, URL|None), ...] 的列表"""
    pattern = r'\[([^\]]+)\]\(([^\)]+)\)'
    segments = []
    last_end = 0
    for match in re.finditer(pattern, text):
        start, end = match.span()
        if start > last_end:
            segments.append((text[last_end:start], None))
        segments.append((match.group(1), match.group(2)))
        last_end = end
    if last_end < len(text):
        segments.append((text[last_end:], None))
    return segments


def parse_draft(filepath):
    """
    解析 Markdown 草稿，回傳結構化資料。
    
    回傳格式:
    {
        'title': '2026.03.06 工作日誌',
        'date_str': '115年03月06日',
        'categories': [
            {
                'name': '行政組事務',
                'items': [
                    {
                        'title': '任務名稱',
                        'progress': '待辦理/已完成',
                        'lines': ['說明：...', ...]
                    },
                    ...
                ]
            },
            ...
        ]
    }
    """
    with open(filepath, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    result = {
        'title': '',
        'date_str': '',
        'categories': []
    }
    
    # 嘗試從檔名推斷日期
    basename = os.path.basename(filepath)
    date_match = re.search(r'(\d{4})-(\d{2})-(\d{2})', basename)
    if date_match:
        y, m, d = int(date_match.group(1)), int(date_match.group(2)), int(date_match.group(3))
        roc_year = y - 1911
        result['date_str'] = f'{roc_year}年{m:02d}月{d:02d}日'
        result['roc_year'] = str(roc_year)
        result['short_date'] = f'{m}/{d}'
    
    current_category = None
    current_item = None
    
    for line in lines:
        text = line.rstrip()
        stripped = text.strip()
        
        if not stripped:
            continue
        
        # 主標題 (# 開頭)
        if stripped.startswith('# ') and not stripped.startswith('## '):
            result['title'] = stripped.replace('# ', '')
            continue
        
        # 分類標題 (## 開頭)
        if stripped.startswith('## ') and not stripped.startswith('### '):
            cat_name = re.sub(r'^## (📌|🤖)?\s*[一二三四五六七八九十]*、?\s*', '', stripped).strip()
            current_category = {'name': cat_name, 'items': []}
            result['categories'].append(current_category)
            current_item = None
            continue
        
        # 子分類標題 (### 開頭) — 視為新的分類
        if stripped.startswith('### ') and not stripped.startswith('#### '):
            cat_name = re.sub(r'^### ', '', stripped).strip()
            current_category = {'name': cat_name, 'items': []}
            result['categories'].append(current_category)
            current_item = None
            continue
        
        # 任務標題 (#### 開頭)
        if stripped.startswith('#### '):
            item_title = re.sub(r'^#### ', '', stripped).strip()
            if item_title.startswith('本項無資料'):
                continue
            current_item = {
                'title': item_title,
                'progress': '待辦理',
                'lines': []
            }
            if current_category:
                current_category['items'].append(current_item)
            continue
        
        # 項目內容行
        if current_item is not None:
            if stripped.startswith('- **') or stripped.startswith('*   ') or stripped.startswith('    *   '):
                # 檢查是否為進度欄位
                if '**進度**' in stripped or '**進度**：' in stripped:
                    progress_match = re.search(r'\*\*進度\*\*[：:]\s*(.+)', stripped)
                    if progress_match:
                        current_item['progress'] = progress_match.group(1).strip()
                else:
                    current_item['lines'].append(stripped)
            elif stripped and not stripped.startswith('#') and not stripped.startswith('---'):
                current_item['lines'].append(stripped)
    
    # 清理空分類
    result['categories'] = [c for c in result['categories'] if c['items']]
    
    return result


# =============================================
# Word 文件產生
# =============================================

def set_cell_font(cell, font_name='標楷體', font_size=12, bold=False):
    """設定儲存格內所有段落的預設字型"""
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = font_name
            run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
            run.font.size = Pt(font_size)
            run.bold = bold


def set_cell_spacing(cell, line_spacing=276):
    """設定儲存格內段落間距"""
    for paragraph in cell.paragraphs:
        pPr = paragraph._element.get_or_add_pPr()
        spacing = OxmlElement('w:spacing')
        spacing.set(qn('w:before'), '0')
        spacing.set(qn('w:after'), '0')
        spacing.set(qn('w:line'), str(line_spacing))
        spacing.set(qn('w:lineRule'), 'auto')
        pPr.append(spacing)


def set_cell_vertical_alignment(cell, alignment='center'):
    """設定儲存格垂直對齊"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    vAlign = OxmlElement('w:vAlign')
    vAlign.set(qn('w:val'), alignment)
    tcPr.append(vAlign)


def set_cell_shading(cell, color_hex):
    """設定儲存格背景顏色"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color_hex)
    cell._tc.get_or_add_tcPr().append(shading_elm)


def add_run_to_paragraph(paragraph, text, font_name='標楷體', font_size=12, bold=False, color=None, underline=False, subscript=False):
    """在段落中新增格式化文字"""
    run = paragraph.add_run(text)
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(font_size)
    run.bold = bold
    if color:
        run.font.color.rgb = color
    if underline:
        run.font.underline = True
    if subscript:
        # 用 XML 直接設定下標
        rPr = run._element.get_or_add_rPr()
        vAlign = OxmlElement('w:vertAlign')
        vAlign.set(qn('w:val'), 'subscript')
        rPr.append(vAlign)
    return run


def add_hyperlink_to_paragraph(paragraph, text, url, font_name='標楷體', font_size=12):
    """在段落中加入可點擊超連結"""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK,
        is_external=True
    )
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    # 藍色
    color_elm = OxmlElement('w:color')
    color_elm.set(qn('w:val'), '0563C1')
    rPr.append(color_elm)

    # 底線
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)

    # 字型
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rPr.append(rFonts)

    # 字型大小
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), str(font_size * 2))
    rPr.append(sz)
    szCs = OxmlElement('w:szCs')
    szCs.set(qn('w:val'), str(font_size * 2))
    rPr.append(szCs)

    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def write_rich_text_to_paragraph(paragraph, text, font_name='標楷體', font_size=12, bold=False):
    """
    將可能含有 Markdown 連結的文字寫入段落。
    純文字部分使用一般 run，連結部分使用 hyperlink。
    """
    segments = parse_markdown_links(text)
    for seg_text, seg_url in segments:
        if seg_url:
            add_hyperlink_to_paragraph(paragraph, seg_text, seg_url, font_name, font_size)
        else:
            add_run_to_paragraph(paragraph, seg_text, font_name, font_size, bold)


def add_category_row(table, text):
    """新增一個合併 3 欄的分類標題列（淺灰底 F2F2F2、14pt 粗體、置中）"""
    row = table.add_row()

    # 🔑 關鍵：從 XML 層級移除多餘的 <w:tc>，只保留第一個
    # table.add_row() 會建立 3 個 <w:tc>，只設 gridSpan 不會自動移除它們
    tr = row._tr
    tcs = tr.findall(qn('w:tc'))
    for tc in tcs[1:]:
        tr.remove(tc)

    # 設定 gridSpan 合併三欄
    cell = row.cells[0]
    tcPr = cell._tc.get_or_add_tcPr()
    gridSpan = OxmlElement('w:gridSpan')
    gridSpan.set(qn('w:val'), '3')
    tcPr.append(gridSpan)

    # 設定文字
    cell.text = text

    # 設定淺灰底色
    set_cell_shading(cell, 'F2F2F2')

    # 設定文字格式
    for para in cell.paragraphs:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in para.runs:
            run.font.name = '標楷體'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')
            run.font.size = Pt(14)
            run.bold = True

    set_cell_spacing(cell)
    return row


def _split_title_and_date(item_title):
    """將項目標題拆分為主文字與日期後綴。
    
    支援的日期後綴格式：
      - ' 115.03.09'（民國年.月.日）
      - ' 2026.03.09'（西元年.月.日）
    
    回傳 (主文字, 日期後綴字串 or None)
    日期後綴字串格式：'-YYYY.MM.DD（星期）'
    """
    # 嘗試匹配尾部的日期（民國或西元）
    match = re.search(r'\s+(\d{2,4})[./](\d{2})[./](\d{2})\s*$', item_title)
    if not match:
        return item_title, None
    
    main_text = item_title[:match.start()]
    year_str = match.group(1)
    month = int(match.group(2))
    day = int(match.group(3))
    
    # 轉換為西元年
    year = int(year_str)
    if year < 200:  # 民國年
        year += 1911
    
    # 計算星期
    weekday_names = ['一', '二', '三', '四', '五', '六', '日']
    try:
        dt = datetime(year, month, day)
        weekday = weekday_names[dt.weekday()]
    except ValueError:
        weekday = '?'
    
    date_suffix = f'-{year}.{month:02d}.{day:02d}（{weekday}）'
    return main_text, date_suffix


def add_data_row(table, item_title, content_lines, progress='待辦理'):
    """新增一筆資料列，含項目、內容、進度"""
    row = table.add_row()

    # ---- 項目欄 ----
    cell_item = row.cells[0]
    p_item = cell_item.paragraphs[0]

    # 拆分主文字與日期後綴，日期以下標呈現（9pt、下標格式 -YYYY.MM.DD（星期），符合 SKILL.md）
    main_text, date_suffix = _split_title_and_date(item_title)
    add_run_to_paragraph(p_item, main_text)
    if date_suffix:
        add_run_to_paragraph(p_item, date_suffix, font_size=9, subscript=True)
    
    set_cell_spacing(cell_item)
    set_cell_vertical_alignment(cell_item, 'top')

    # ---- 內容欄（多層自動大綱編號：一、（一）、1.、（1）、a.、（a））----
    cell_content = row.cells[1]
    first_para = True
    outline_level = 0  # 追蹤目前大綱層級

    # 重設編號，讓每個項目都從「一、」重新開始
    # 每個 cell 都需要獨立的 numId，這樣 level 0 才會從 1 重新計數
    _new_outline_num_id()  # 建立新的 numId，level 0 startOverride=1

    for line in content_lines:
        stripped = line.strip()

        # 計算縮排層級
        indent_level = 0
        if line.startswith('    *'):
            indent_level = 2  # 第三層以下
        elif line.startswith('*'):
            indent_level = 1  # 第二層

        # 移除 Markdown 前綴
        if stripped.startswith('- **'):
            content = stripped[2:]
            label_match = re.match(r'\*\*([^*]+)\*\*[：:]\s*(.*)', content)
            if label_match:
                label = label_match.group(1)
                value = label_match.group(2)

                if label == '說明':
                    # 第一層：一、
                    if first_para:
                        p = cell_content.paragraphs[0]
                        first_para = False
                    else:
                        p = cell_content.add_paragraph()
                    apply_numbering_to_paragraph(p, 0)  # 一、
                    write_rich_text_to_paragraph(p, value)
                    outline_level = 0
                elif label == '附件':
                    if first_para:
                        p = cell_content.paragraphs[0]
                        first_para = False
                    else:
                        p = cell_content.add_paragraph()
                    write_rich_text_to_paragraph(p, f'附件：{value}')
                else:
                    if first_para:
                        p = cell_content.paragraphs[0]
                        first_para = False
                    else:
                        p = cell_content.add_paragraph()
                    apply_numbering_to_paragraph(p, outline_level)
                    cleaned = content.replace('**', '')
                    write_rich_text_to_paragraph(p, cleaned)

        elif stripped.startswith('*   **') or stripped.startswith('    *   **'):
            # 第二層：（一）或第三層 1.
            content = stripped.lstrip('* ').lstrip()
            label_match = re.match(r'\*\*([^*]+)\*\*[：:]\s*(.*)', content)

            if first_para:
                p = cell_content.paragraphs[0]
                first_para = False
            else:
                p = cell_content.add_paragraph()

            if indent_level >= 2:
                apply_numbering_to_paragraph(p, 2)  # 1.
            else:
                apply_numbering_to_paragraph(p, 1)  # （一）
                outline_level = 1

            if label_match:
                label = label_match.group(1)
                value = label_match.group(2)
                add_run_to_paragraph(p, f'{label}：', bold=True)
                write_rich_text_to_paragraph(p, value)
            else:
                cleaned = content.replace('**', '')
                write_rich_text_to_paragraph(p, cleaned)

        elif stripped.startswith('*   ') or stripped.startswith('    *   '):
            content = stripped.lstrip('* ').lstrip()
            if first_para:
                p = cell_content.paragraphs[0]
                first_para = False
            else:
                p = cell_content.add_paragraph()

            if indent_level >= 2:
                apply_numbering_to_paragraph(p, 2)  # 1.
            else:
                apply_numbering_to_paragraph(p, 1)  # （一）
                outline_level = 1

            write_rich_text_to_paragraph(p, content)

        else:
            if first_para:
                p = cell_content.paragraphs[0]
                first_para = False
            else:
                p = cell_content.add_paragraph()
            apply_numbering_to_paragraph(p, outline_level if outline_level > 0 else 0)
            write_rich_text_to_paragraph(p, stripped)

    set_cell_spacing(cell_content)
    set_cell_vertical_alignment(cell_content, 'top')

    # ---- 進度欄 ----
    cell_progress = row.cells[2]
    p_progress = cell_progress.paragraphs[0]
    p_progress.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run_to_paragraph(p_progress, progress)
    set_cell_spacing(cell_progress)
    set_cell_vertical_alignment(cell_progress, 'center')

    return row


def generate_docx(data, output_path):
    """產生 Word 工作日誌文件"""
    doc = Document()

    # === 初始化多層大綱編號引擎 ===
    init_outline_numbering(doc)

    # === 設定預設字型 ===
    style = doc.styles['Normal']
    font = style.font
    font.name = '標楷體'
    font.size = Pt(12)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')
    style.paragraph_format.line_spacing = 1.15

    # === 頁面邊距 ===
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    # === 文件標題 ===
    p_org = doc.add_paragraph()
    p_org.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_org = p_org.add_run('財團法人台中市私立慈光社會福利慈善事業基金會')
    run_org.font.name = '標楷體'
    run_org._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')
    run_org.font.size = Pt(16)
    run_org.bold = True

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run('工作日誌')
    run_title.font.name = '標楷體'
    run_title._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')
    run_title.font.size = Pt(16)
    run_title.bold = True

    # === 資訊行（填表人靠右） ===
    p_info = doc.add_paragraph()
    roc_year = data.get('roc_year', '115')
    date_str = data.get('date_str', '')

    # 左側資訊
    run1 = p_info.add_run(f'年度：{roc_year}年　　日期：{date_str}')
    run1.font.name = '標楷體'
    run1._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')
    run1.font.size = Pt(12)
    run1.bold = True

    # 使用定位點將填表人靠右
    # 設定 right indent 讓文字從右側縮回，並用 tab 分隔
    p_pr = p_info._p.get_or_add_pPr()

    # 設定 right indent（設為負值讓內容靠右）
    ind = OxmlElement('w:ind')
    ind.set(qn('w:right'), '4320')  # 3cm right indent
    p_pr.append(ind)

    # 建立 tab stops
    tabs = OxmlElement('w:tabs')
    # 左側 tab（在 0 位置，幫助年度保持在左）
    left_tab = OxmlElement('w:tab')
    left_tab.set(qn('w:val'), 'left')
    left_tab.set(qn('w:pos'), '0')
    tabs.append(left_tab)
    # 右側 tab（讓填表人靠右）
    right_tab = OxmlElement('w:tab')
    right_tab.set(qn('w:val'), 'right')
    right_tab.set(qn('w:pos'), '8640')  # 6cm from left = 表格右側附近
    tabs.append(right_tab)
    p_pr.append(tabs)

    run2 = p_info.add_run('\t填表人：李冠葦')
    run2.font.name = '標楷體'
    run2._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')
    run2.font.size = Pt(12)
    run2.bold = True

    # === 建立表格（3欄）===
    table = doc.add_table(rows=1, cols=3, style='Table Grid')
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False

    # 設定欄寬（項目 4.0cm、內容 9.5cm、進度 2.5cm = 16.0cm）
    # 注意：A4 頁面扣除左右邊距後可用約 16.5cm，不可超過
    col_widths = [Cm(4.0), Cm(9.5), Cm(2.5)]
    for i, width in enumerate(col_widths):
        table.columns[i].width = width
        table.rows[0].cells[i].width = width

    # 表頭（灰底 D9D9D9）
    headers = ['項目', '內容', '進度']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run_to_paragraph(p, header, bold=True)
        
        # 設定灰底
        set_cell_shading(cell, 'D9D9D9')
        
        set_cell_spacing(cell)
        set_cell_vertical_alignment(cell, 'center')

    # === 填入分類與資料 ===
    for category in data['categories']:
        add_category_row(table, category['name'])

        for item in category['items']:
            progress = item.get('progress', '待辦理')
            
            # 從內容推斷進度（如果沒有明確指定）
            if progress == '待辦理':
                for line in item.get('lines', []):
                    if '已完成' in line or '已辦理' in line or '已完成' in line:
                        progress = '已完成'
                        break
                    elif '持續' in line or '進行中' in line:
                        progress = '持續追蹤中'
                        break

            add_data_row(
                table,
                item_title=item['title'],
                content_lines=item.get('lines', []),
                progress=progress
            )

    # === 儲存 ===
    doc.save(output_path)
    print(f'✅ 工作日誌已產出：{output_path}')


# =============================================
# 主程式
# =============================================

def main():
    if len(sys.argv) < 2:
        print('使用方式: python3 gen_work_log_table.py <草稿路徑> [輸出路徑]')
        sys.exit(1)

    draft_path = os.path.expanduser(sys.argv[1])

    if len(sys.argv) >= 3:
        output_path = os.path.expanduser(sys.argv[2])
    else:
        # 從草稿檔名推斷輸出檔名
        basename = os.path.basename(draft_path)
        date_match = re.search(r'(\d{4})-(\d{2})-(\d{2})', basename)
        if date_match:
            y, m, d = date_match.group(1), date_match.group(2), date_match.group(3)
            output_name = f'{y}.{m}.{d}_工作日誌.docx'
        else:
            output_name = '工作日誌.docx'
        # 預設輸出到工作日誌資料夾
        output_path = os.path.join(os.path.expanduser('~/個人app/個人資料夾/工作日誌'), output_name)

    if not os.path.exists(draft_path):
        print(f'❌ 找不到草稿檔案：{draft_path}')
        sys.exit(1)

    # 確保輸出目錄存在
    output_dir = os.path.dirname(output_path)
    if output_dir and not os.path.exists(output_dir):
        os.makedirs(output_dir)

    data = parse_draft(draft_path)
    generate_docx(data, output_path)


if __name__ == '__main__':
    main()
